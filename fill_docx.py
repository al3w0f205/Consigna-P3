import docx
import sys
import os
from docx.oxml import OxmlElement

def is_file_locked(file_path):
    """
    Verifica si un archivo esta bloqueado por otro proceso (por ejemplo, Microsoft Word)
    intentando abrirlo en modo de lectura/escritura exclusiva.
    """
    if not os.path.exists(file_path):
        return False
    try:
        # Intentamos abrir el archivo con acceso exclusivo
        with open(file_path, "r+"):
            pass
        return False
    except (IOError, PermissionError):
        return True

def insert_paragraph_after(paragraph, text, is_header=False):
    p_element = paragraph._element
    new_p_element = OxmlElement('w:p')
    p_element.addnext(new_p_element)
    new_paragraph = docx.text.paragraph.Paragraph(new_p_element, paragraph._parent)
    
    # Copiar estilo e interlineados del parrafo ancla
    new_paragraph.style = paragraph.style
    if paragraph.paragraph_format:
        new_paragraph.paragraph_format.space_before = paragraph.paragraph_format.space_before
        new_paragraph.paragraph_format.space_after = paragraph.paragraph_format.space_after
        new_paragraph.paragraph_format.line_spacing = paragraph.paragraph_format.line_spacing
        new_paragraph.paragraph_format.alignment = paragraph.paragraph_format.alignment
        
    font_name = None
    font_size = None
    font_color = None
    
    if paragraph.runs:
        r_base = paragraph.runs[0]
        if r_base.font:
            font_name = r_base.font.name
            font_size = r_base.font.size
            if r_base.font.color and r_base.font.color.rgb:
                font_color = r_base.font.color.rgb
                
    run = new_paragraph.add_run(text)
    if is_header:
        run.bold = True
        
    if font_name:
        run.font.name = font_name
    if font_size:
        run.font.size = font_size
    if font_color:
        run.font.color.rgb = font_color
        
    return new_paragraph

def set_cell_text_formatted(cell, text):
    p_base = cell.paragraphs[0]
    font_name = None
    font_size = None
    font_color = None
    
    # Buscar formato existente en los parrafos de la celda
    for p in cell.paragraphs:
        if p.runs:
            r_base = p.runs[0]
            if r_base.font:
                font_name = r_base.font.name
                font_size = r_base.font.size
                if r_base.font.color and r_base.font.color.rgb:
                    font_color = r_base.font.color.rgb
                break
                
    # Limpiar parrafos adicionales si los hubiera
    while len(cell.paragraphs) > 1:
        p_extra = cell.paragraphs[-1]._element
        p_extra.getparent().remove(p_extra)
        
    p_base.text = ""
    run = p_base.add_run(text)
    if font_name:
        run.font.name = font_name
    if font_size:
        run.font.size = font_size
    if font_color:
        run.font.color.rgb = font_color

def clean_previous_insertions(doc):
    paragraphs_to_remove = []
    keywords = [
        "### Alternativa 1", "Esta alternativa plantea", "El historial de los ultimos 30",
        "### Alternativa 2", "Esta alternativa propone", "### Analisis Comparativo",
        "Alternativa 1 - Ventajas", "Alternativa 1 - Desventajas", "Alternativa 2 - Ventajas", "Alternativa 2 - Desventajas",
        "La Alternativa 2 es", "Al evaluar y elegir una", "### Fundamentacion Tecnica",
        "### Garantia de Eficiencia", "### Estructura de la Solucion", "Enlace al Repositorio",
        "Estructuras de Datos Definidas", "- 'typedef struct", "La persistencia se implementa",
        "Nota de Diagramacion", "### Casos de Prueba", "Caso de Prueba 1", "Caso de Prueba 2",
        "Caso de Prueba 3", "### INTRODUCION", "### OBJETIVOS DEL TRABAJO", "### CONCLUSIONES",
        "### RECOMENDACIONES", "### REFERENCIAS BIBLIOGRAFICAS", "La contaminacion del aire en",
        "Diseñar, desarrollar y validar", "Dimension Tecnologica", "Dimension Ambiental", "Dimension Socioeconomica"
    ]
    
    for p in doc.paragraphs:
        p_text = p.text.strip()
        for kw in keywords:
            if kw in p_text:
                paragraphs_to_remove.append(p)
                break
                
    print(f"Borrando {len(paragraphs_to_remove)} parrafos previos...")
    for p in paragraphs_to_remove:
        p_element = p._element
        p_element.getparent().remove(p_element)

def fill_document(docx_path):
    # Validacion previa de bloqueo de archivo
    if is_file_locked(docx_path):
        print(f"\n[ERROR] El archivo Word esta abierto en otro programa (como Microsoft Word).")
        print(f"Por favor, cierre el documento '{os.path.basename(docx_path)}' y vuelva a intentarlo.\n")
        sys.exit(1)

    print(f"Abriendo documento: {docx_path}")
    doc = docx.Document(docx_path)
    
    # Limpiar inserciones de parrafos previas
    clean_previous_insertions(doc)
    
    # 1. Rellenar Tabla 0 (Formulacion del Problema)
    print("Rellenando la Tabla de Formulacion del Problema (Tabla 0) con formato heredado...")
    t0 = doc.tables[0]
    
    set_cell_text_formatted(t0.cell(0, 1), "Sistema de Gestion y Prediccion de Calidad del Aire en Zonas Urbanas (SIGPA-Quito)")
    set_cell_text_formatted(t0.cell(1, 1), "Autoridades municipales de Quito (ambiente y salud) y analistas de datos meteorologicos.")
    set_cell_text_formatted(t0.cell(2, 1), "Monitorear contaminantes urbanos (CO2, SO2, NO2, PM2.5), calcular predicciones a 24 horas integrando factores climaticos y almacenar de forma persistente alertas y reportes.")
    set_cell_text_formatted(t0.cell(3, 1), (
        "1. Nombre de zona y mediciones de contaminacion (CO2, SO2, NO2, PM2.5).\n"
        "2. Datos climaticos (temperatura, viento, humedad).\n"
        "3. Configuracion de limites OMS y tope maximo de zonas.\n"
        "4. Historial de los ultimos 30 dias."
    ))
    set_cell_text_formatted(t0.cell(4, 1), (
        "1. Diagnostico de calidad del aire actual (Normal/Excedido).\n"
        "2. Prediccion a 24 horas y alertas preventivas futuras.\n"
        "3. Reporte general estructurado ('reporte_calidad_aire.txt') y bitacora de log de excesos ('alertas.log')."
    ))
    set_cell_text_formatted(t0.cell(5, 1), (
        "1. Carga y guardado binario estructurado en 'datos.dat'.\n"
        "2. Gestion dinamica del arreglo de zonas con 'realloc'.\n"
        "3. Administracion del historial mediante cola circular de tamaño 30 en O(1).\n"
        "4. Calculo de promedio ponderado y factor climatico predictivo."
    ))
    set_cell_text_formatted(t0.cell(6, 1), (
        "1. Compilacion limpia bajo estandar C99 sin librerias externas.\n"
        "2. Limites rigurosos segun directrices OMS.\n"
        "3. Registro cronologico automatico del sistema en alertas.log para auditorias."
    ))
    set_cell_text_formatted(t0.cell(7, 1), (
        "1. Historial estatico acotado a un maximo de 30 dias en memoria.\n"
        "2. Arreglo maximo configurable hasta 100 zonas.\n"
        "3. Persistencia local en formato binario e interfaz interactiva por consola CLI."
    ))
    
    # 2. Secciones de Texto Sintetizadas
    sections_to_add = [
        {
            "anchor": "Plantea varias alternativas de soluci\u00f3n",
            "content": [
                ("### Alternativa 1: Arquitectura Dinamica y Cola Circular (Solucion Seleccionada)", True),
                ("Usa un arreglo dinamico de estructuras de tipo 'Zona' redimensionado con 'realloc' para optimizar la RAM. El historial de 30 dias se gestiona mediante una cola circular de tamaño estatico administrada con indices de control ('head' y 'count') con coste O(1). La persistencia es binaria rapida en 'datos.dat' mediante 'fwrite' y 'fread'.", False),
                ("### Alternativa 2: Arreglos Paralelos Estaticos y Archivos de Texto", True),
                ("Usa multiples arreglos planos estaticos de tamaño fijo (ej. nombres, contaminantes) definidos en compilacion. El historial de 30 dias se almacena en una matriz bidimensional, requiriendo desplazar fisicamente la memoria en cada insercion diaria (coste O(N)). La persistencia usa archivos CSV de texto con conversion manual mediante 'fscanf' y 'fprintf'.", False)
            ]
        },
        {
            "anchor": "Analiza las alternativas anteriormente planteadas",
            "content": [
                ("### Analisis Comparativo de Ventajas y Desventajas", True),
                ("- Alternativa 1: Ventajas en optimizacion de RAM, insercion rapida en O(1) y almacenamiento binario directo. Desventaja en complejidad de depuracion de punteros y persistencia binaria.", False),
                ("- Alternativa 2: Ventajas en simplicidad del codigo al evitar punteros y legibilidad directa del archivo de texto. Desventaja en desperdicio de RAM y lentitud de I/O por conversion de archivos de texto.", False)
            ]
        },
        {
            "anchor": "\u00bfCu\u00e1l de estas alternativas ser\u00eda la m\u00e1s f\u00e1cil de implementar y por qu\u00e9?",
            "content": [
                ("La Alternativa 2 es la mas facil de implementar al prescindir de memoria dinamica, punteros dobles y aritmetica de colas circulares.", False)
            ]
        },
        {
            "anchor": "\u00bfQu\u00e9 aspectos son importantes al elegir una soluci\u00f3n eficiente?",
            "content": [
                ("Los aspectos clave para una solucion eficiente son la complejidad algoritmica en tiempo, la optimizacion del uso de RAM, el coste de acceso a disco (I/O) y la mantenibilidad modular.", False)
            ]
        },
        {
            "anchor": "Seleccione la mejor alternativa para ser desarrollada, en base al an\u00e1lisis",
            "content": [
                ("### Fundamentacion Tecnica de la Seleccion de la Solucion", True),
                ("Se selecciona la Alternativa 1. En monitoreo continuo de sensores, un coste constante O(1) de insercion es imprescindible para prevenir latencias en la adquisicion de datos. La encapsulacion en estructuras dinamicas y persistencia binaria garantizan un sistema robusto y escalable.", False),
                ("### Garantia de Eficiencia y Claridad en el Codigo", True),
                ("Garantizaria la eficiencia mediante analisis asintotico Big-O y herramientas de perfilado de CPU (profiler). La claridad se asegura con programacion modular, nombres de variables autodocumentados y comentarios claros.", False)
            ]
        },
        {
            "anchor": "Desarrolle la alternativa seleccionada, haciendo uso de diagramas",
            "content": [
                ("### Estructura de la Solucion e Implementacion en C", True),
                ("El programa se estructura en tres modulos principales: 'funciones.h' (cabeceras y tipos), 'funciones.c' (algoritmos y persistencia) y 'main.c' (menu y bucle). Define estructuras como Clima, Contaminacion, RegistroHistorial y Zona. Si no existe base binaria previa, se inicializan automaticamente 5 zonas piloto de Quito: Centro Historico, Belisario, Carapungo, El Camal y Guamani.", False),
                ("Enlace al Repositorio de GitHub: https://github.com/carlosguaita/calidad-aire-quito", False)
            ]
        },
        {
            "anchor": "Valide la soluci\u00f3n desarrollada funcione correctamente",
            "content": [
                ("### Validacion con Casos de Prueba", True),
                ("- Caso 1: Al eliminar 'datos.dat', el sistema detecta la ausencia, inicializa las 5 zonas de Quito por defecto en memoria y crea una nueva base de datos binaria. Exito verificado.", False),
                ("- Caso 2: Al registrar una zona con contaminacion excedida, el sistema la añade dinamicamente y registra de forma automatica la anomalia en 'alertas.log' con la marca de tiempo de la maquina. Exito verificado.", False),
                ("- Caso 3: Al exportar reporte, el programa calcula las predicciones a 24 horas integrando temperatura, viento e humedad y genera el archivo plano 'reporte_calidad_aire.txt' con recomendaciones. Exito verificado.", False)
            ]
        },
        {
            "anchor": "Toda afirmaci\u00f3n y justificaci\u00f3n dentro del informe deber\u00e1 contener referencias",
            "content": [
                ("### INTRODUCION", True),
                ("Quito afronta severos problemas de contaminacion del aire (CO2, SO2, NO2, PM2.5) debido al parque automotor y su topografia, afectando la salud publica. Se requiere un sistema de software modular en C para monitoreo persistente y emision de alertas predictivas basadas en promedios historicos y factores climaticos.", False),
                ("### OBJETIVOS DEL TRABAJO", True),
                ("Objetivo General: Desarrollar una herramienta de software modular en C que gestione de forma persistente la calidad del aire y prediga la contaminacion a 24 horas usando variables meteorologicas.", False),
                ("Objetivos Especificos:\n1. Implementar memoria dinamica y colas circulares para un historial eficiente de 30 dias en O(1).\n2. Programar un algoritmo de prediccion ponderada con dispersion meteorologica.\n3. Automatizar persistencia binaria y logs de auditoria.", False),
                ("### CONCLUSIONES", True),
                ("1. Las colas circulares y arreglos dinamicos optimizan el uso de memoria y CPU, ideales para entornos embebidos de medicion.\n2. Integrar factores meteorologicos (viento, temperatura, humedad) mejora sustancialmente el realismo del modelo predictivo.\n3. La persistencia binaria directa disminuye la latencia de disco y previene la corrupcion de datos.", False),
                ("### RECOMENDACIONES", True),
                ("1. Conectar fisicamente sensores para automatizar la captura de datos en tiempo real.\n2. Implementar tecnicas de series temporales para predicciones complejas a largo plazo.", False),
                ("### REFERENCIAS BIBLIOGRAFICAS (APA 7)", True),
                ("- Organizacion Mundial de la Salud (OMS). (2021). Directrices mundiales sobre la calidad del aire. OMS.\n- Secretaria de Ambiente de Quito. (2022). Informe Anual de Calidad del Aire. Distrito Metropolitano de Quito.", False)
            ]
        }
    ]
    
    # 3. Aplicar las secciones
    for sec in sections_to_add:
        anchor_text = sec["anchor"]
        content = sec["content"]
        
        found_p = None
        for p in doc.paragraphs:
            if anchor_text in p.text:
                found_p = p
                break
                
        if found_p:
            print(f"Insertando contenido para el ancla: '{anchor_text}'")
            # Invertimos para insertar en el orden correcto
            for text, is_header in reversed(content):
                insert_paragraph_after(found_p, text, is_header)
        else:
            print(f"[!] No se encontro el ancla: '{anchor_text}'")
            
    print("Guardando documento modificado...")
    doc.save(docx_path)
    print("Documento guardado con exito.")

if __name__ == "__main__":
    fill_document(r"C:\Users\power\Desktop\ActividadEvaluativa_RC2_ ISWZ1102_202610.docx")
